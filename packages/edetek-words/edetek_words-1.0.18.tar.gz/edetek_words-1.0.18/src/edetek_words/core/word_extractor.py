# !/usr/bin/python3
# -*- coding:utf-8 -*-
"""
@Author: xiaodong.li
@Time: 7/28/2025 4:53 PM
@Description: Description
@File: word_extractor.py
"""
import re
from pathlib import Path
from typing import List, Set, Optional

from docx import Document

from edetek_words.common.docx_utils import get_bold_semantic_texts
from edetek_words.common.json_utils import save_json
from edetek_words.common.path import output_doc_path
from edetek_words.common.word_cleaner import clean_data
from edetek_words.dto.segment_dto import SegmentDTO
from edetek_words.dto.styled_text_segment import StyledTextSegment, FontStyle


def exclude_text(text):
    if re.fullmatch(r"\d+", text):
        return True
    if re.fullmatch(r"\d{1,2}:\d{2}", text):
        return True
    return False


class WordExtractor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.segments: List[SegmentDTO] = []
        self._seen: Set[str] = set()

    def extract(self) -> List[SegmentDTO]:
        self._process_paragraphs(self.doc.paragraphs)
        self._process_tables(self.doc.tables)
        self._process_sections()
        filename = Path(self.file_path).with_suffix(".json").name
        filepath = output_doc_path(filename)
        if filepath.exists() and filepath.is_file():
            filepath.unlink()
        save_json(str(filepath), [segment.original_text for segment in self.segments], ensure_ascii=False)  # for debug
        return self.segments

    def _process_paragraphs(self, paragraphs):
        for para in paragraphs:
            self._handle_paragraph(para)

    def _process_tables(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self._process_paragraphs(cell.paragraphs)

    def _process_sections(self):
        for section in self.doc.sections:
            self._process_paragraphs(section.header.paragraphs)
            self._process_tables(section.header.tables)
            self._process_paragraphs(section.footer.paragraphs)

    def _handle_paragraph(self, para):
        full_text = para.text
        cleaned_text = clean_data(full_text)
        if not cleaned_text:
            return
        for line in re.split(r'[\n\t]+', cleaned_text):
            cleaned_line = clean_data(line)
            if not cleaned_line or cleaned_line in self._seen:
                return
            if exclude_text(cleaned_line):
                return
            segment = SegmentDTO(
                original_text=cleaned_line,
            )
            self.segments.append(segment)
            self._seen.add(cleaned_line)
        texts = get_bold_semantic_texts(para)
        if texts:
            segment: SegmentDTO = self.get_segment_by_original_text(cleaned_text)
            segment.styled_segments = [StyledTextSegment(text, [FontStyle.BOLD]) for text in texts]

    def get_segment_by_original_text(self, original_text) -> Optional[SegmentDTO]:
        for segment in self.segments:
            if segment.original_text == original_text:
                return segment
        return None
