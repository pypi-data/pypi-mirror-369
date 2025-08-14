#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
基于LDA主题模型的文档分段工具

这个脚本使用LDA主题模型识别文本中的主题变化点，并据此进行分段。
支持处理单个Word文档或包含多个Word文档的文件夹。
"""

import os
import re
import argparse
import logging
import jieba
import numpy as np
from docx import Document
from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import nltk
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.metrics.pairwise import cosine_similarity

# 确保nltk数据已下载


# 配置日志
logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s', level=logging.INFO)

# 中文停用词列表
CHINESE_STOP_WORDS = set([
    '的', '了', '和', '是', '就', '都', '而', '及', '与', '这', '那', '有', '在', '中', '为',
    '对', '地', '也', '得', '着', '之', '于', '上', '下', '但', '因', '此', '没', '很', '只',
    '些', '如', '它', '们', '我', '你', '他', '她', '或', '其', '可', '能', '要', '自', '以',
    '却', '被', '所', '然', '已', '还', '从', '到', '呢', '啊', '吧', '吗', '呀', '哦', '哪',
    '么', '什', '怎', '啥', '哎', '嗯', '哈', '啥', '哇', '嘿', '哟', '嘛', '咋', '哩', '咱',
    '哼', '唉', '嗨', '呐', '呵', '嘻', '嗬', '咦', '嘎', '呦', '嘭', '咚', '咔', '嚓', '啪',
    '啧', '呸', '呲', '咝', '哧', '嘶', '嗖', '嗒', '嗝', '嗡', '嗷', '嗯', '嗬', '嗲', '嗳',
    '嗨', '嗯', '嗡', '嗷', '嗝', '嗨', '嗬', '嗲', '嗳', '嗯', '嗡', '嗷', '嗝', '嗨', '嗬',
    '嗲', '嗳'
])


class LDASegmenter:
    """基于LDA主题模型的文档分段器"""

    def __init__(self,min_paragraph_length=150, max_paragraph_length=400,
                 num_topics=8, similarity_threshold=0.35, verbose=False, nltk_data_path=None):
        """
        初始化分段器

        参数:
            min_paragraph_length (int): 最小段落长度（字符数）
            max_paragraph_length (int): 最大段落长度（字符数）
            num_topics (int): LDA主题模型的主题数量
            similarity_threshold (float): 主题相似度阈值，低于此值视为主题变化
            verbose (bool): 是否输出详细信息
        """
        if nltk_data_path:
            nltk.data.path.append(nltk_data_path)
        try:
            nltk.data.find('tokenizers/punkt')
            print(nltk.data.find('tokenizers/punkt'))
            nltk.data.find('corpora/stopwords')
            print(nltk.data.find('tokenizers/punkt'))
        except LookupError:
            nltk.download('punkt')
            nltk.download('stopwords')
        self.min_paragraph_length = min_paragraph_length
        self.max_paragraph_length = max_paragraph_length
        self.num_topics = num_topics
        self.similarity_threshold = similarity_threshold
        self.verbose = verbose

        # 加载英文停用词
        self.stop_words = set(nltk.corpus.stopwords.words('english'))

        # 添加中文停用词
        self.stop_words.update(CHINESE_STOP_WORDS)

        if verbose:
            print(f"初始化LDA分段器: 最小段落长度={min_paragraph_length}, 最大段落长度={max_paragraph_length}, "
                  f"主题数量={num_topics}, 相似度阈值={similarity_threshold}")

    def preprocess_text(self, text):
        """
        预处理文本，包括分句、分词和去除停用词

        参数:
            text (str): 输入文本

        返回:
            list: 预处理后的句子列表，每个句子是一个词语列表
        """
        # 分句
        sentences = []
        for paragraph in text.split('\n'):
            if paragraph.strip():
                # 使用正则表达式分句，处理中英文混合文本
                for sent in re.split(r'(?<=[.。!！?？;；])\s*', paragraph):
                    if sent.strip():
                        sentences.append(sent.strip())

        # 分词和去除停用词
        tokenized_sentences = []
        processed_sentences = []

        for sentence in sentences:
            # 检测语言
            if any('\u4e00' <= char <= '\u9fff' for char in sentence):
                # 中文分词
                words = jieba.cut(sentence)
                #words = [word for word in words if word not in self.stop_words and len(word.strip()) > 1]
            else:
                # 英文分词
                words = word_tokenize(sentence.lower())
                #words = [word for word in words if word not in self.stop_words and word.isalpha()]

            if words:
                tokenized_sentences.append(words)
                processed_sentences.append(' '.join(words))  # 用于CountVectorizer

        return tokenized_sentences, processed_sentences

    def build_lda_model(self, processed_sentences):
        """
        构建LDA主题模型

        参数:
            processed_sentences (list): 预处理后的句子列表（空格分隔的词语字符串）

        返回:
            tuple: (CountVectorizer, LDA模型, 文档-主题矩阵)
        """
        # 如果句子数量太少，无法构建有效的LDA模型
        if len(processed_sentences) < 10:
            # 为每个句子分配随机主题
            if self.verbose:
                print(f"句子数量太少 ({len(processed_sentences)}), 使用简单分段方法")

            # 创建一个简单的主题分布矩阵
            doc_topic_matrix = np.random.rand(len(processed_sentences), self.num_topics)
            # 归一化，使每行和为1
            row_sums = doc_topic_matrix.sum(axis=1)
            doc_topic_matrix = doc_topic_matrix / row_sums[:, np.newaxis]

            # 返回空的向量化器和模型，但提供有效的主题矩阵
            return None, None, doc_topic_matrix

        # 使用CountVectorizer将文本转换为词频矩阵
        vectorizer = CountVectorizer(min_df=1, max_df=0.9)

        try:
            X = vectorizer.fit_transform(processed_sentences)

            # 训练LDA模型
            lda = LatentDirichletAllocation(
                n_components=self.num_topics,
                max_iter=10,
                learning_method='online',
                random_state=0
            )

            lda.fit(X)

            # 获取文档-主题分布
            doc_topic_matrix = lda.transform(X)

            if self.verbose:
                # 打印主题词
                feature_names = vectorizer.get_feature_names_out()
                self._print_top_words(lda, feature_names, 5)

            return vectorizer, lda, doc_topic_matrix

        except ValueError as e:
            # 处理"no terms remain"错误
            if "no terms remain" in str(e) or "empty vocabulary" in str(e):
                if self.verbose:
                    print(f"无法构建LDA模型: {e}，使用简单分段方法")

                # 创建一个简单的主题分布矩阵
                doc_topic_matrix = np.random.rand(len(processed_sentences), self.num_topics)
                # 归一化，使每行和为1
                row_sums = doc_topic_matrix.sum(axis=1)
                doc_topic_matrix = doc_topic_matrix / row_sums[:, np.newaxis]

                return None, None, doc_topic_matrix
            else:
                raise e

    def _print_top_words(self, model, feature_names, n_top_words):
        """
        打印每个主题的前N个关键词

        参数:
            model: LDA模型
            feature_names: 特征名称（词汇表）
            n_top_words: 每个主题要显示的关键词数量
        """
        print(f"LDA模型训练完成，主题数量: {self.num_topics}")
        for topic_idx, topic in enumerate(model.components_):
            top_words_idx = topic.argsort()[:-n_top_words - 1:-1]  # 获取前N个词的索引
            top_words = [feature_names[i] for i in top_words_idx]
            print(f"主题 {topic_idx}: {' '.join(top_words)}")

    def detect_topic_shifts(self, doc_topic_matrix):
        """
        检测主题变化点

        参数:
            doc_topic_matrix (numpy.ndarray): 文档-主题矩阵

        返回:
            list: 主题变化点的索引列表
        """
        topic_shifts = []

        # 计算相邻句子的主题相似度，检测主题变化点
        for i in range(1, len(doc_topic_matrix)):
            # 使用余弦相似度计算主题分布的相似度
            similarity = cosine_similarity([doc_topic_matrix[i - 1]], [doc_topic_matrix[i]])[0][0]

            # 如果相似度低于阈值，认为发生了主题变化
            if similarity < self.similarity_threshold:
                topic_shifts.append(i)
                if self.verbose:
                    print(f"检测到主题变化点: 句子 {i}, 相似度: {similarity:.4f}")

        return topic_shifts

    def _segment_text_by_topic(self, sentences, sentence_topics):
        """
        根据主题变化对文本进行分段

        参数:
            sentences (list): 句子列表
            sentence_topics (list): 每个句子的主题分布

        返回:
            list: 分段后的文本列表
        """
        segments = []
        current_segment = []
        prev_topic = None

        # 检测段落标记词
        paragraph_markers = ["第一", "第二", "第三", "第四", "第五", "首先", "其次", "再次", "然后", "最后", "总之",
                             "总结", "最终", "一方面", "另一方面", "因此", "所以", "然而", "但是", "不过", "另外",
                             "此外", "除此之外", "总的来说", "综上所述"]

        # 说话人标记的正则表达式
        speaker_pattern = re.compile(r'(^|\n)([A-Za-z\u4e00-\u9fa5]+[：:].{0,10})')

        for i, (sentence, topic) in enumerate(zip(sentences, sentence_topics)):
            # 检查是否是段落的自然开始
            is_new_paragraph = False

            # 检查是否有说话人变化（如"张三："或"李四:"）
            speaker_match = speaker_pattern.search(sentence)
            if speaker_match and i > 0:
                is_new_paragraph = True

            # 检查句子是否包含段落标记词
            for marker in paragraph_markers:
                if marker in sentence and i > 0:  # 不处理第一个句子
                    # 确保标记词在句子开头位置（前10个字符内）
                    if sentence.find(marker) < 10 and len(current_segment) > 0:
                        is_new_paragraph = True
                        break

            # 检查是否是自然停顿点或逻辑转折处
            if i > 0 and len(current_segment) > 0:
                # 检查是否是问句结尾（通常表示话题转换）
                if sentences[i - 1].rstrip().endswith(("?", "？")) and not sentence.startswith(
                        ("是的", "对", "没错", "不是", "不", "可能")):
                    is_new_paragraph = True

                # 检查是否是明显的逻辑转折
                if sentence.lstrip().startswith(("但", "但是", "然而", "不过", "可是", "相反", "反之")):
                    is_new_paragraph = True

            # 计算与前一句的主题相似度
            if prev_topic is not None:
                similarity = np.dot(prev_topic, topic)
                if self.verbose:
                    print(f"检测到主题变化点: 句子 {i}, 相似度: {similarity:.4f}")

                # 主题变化显著，考虑分段
                if similarity < self.similarity_threshold:
                    # 只有当当前段落长度达到一定程度时才考虑分段
                    current_segment_text = "".join(current_segment)
                    if len(current_segment_text) >= self.min_paragraph_length:
                        is_new_paragraph = True

            # 如果是新段落的开始，且当前段落不为空，则结束当前段落
            if is_new_paragraph and current_segment:
                segments.append("".join(current_segment))
                current_segment = []

            # 添加当前句子到当前段落
            current_segment.append(sentence)
            prev_topic = topic

        # 添加最后一个段落
        if current_segment:
            segments.append("".join(current_segment))

        # 合并过短的段落并分割过长的段落
        merged_segments = self._merge_short_segments(segments)

        return merged_segments

    def _is_paragraph_start(self, sentence):
        """
        判断句子是否是段落的自然开始

        参数:
            sentence (str): 句子文本

        返回:
            bool: 是否是段落开始
        """
        # 段落开始的标志词和短语
        paragraph_starters = [
            "首先", "其次", "再次", "然后", "最后", "总之", "因此", "所以", "然而", "但是",
            "不过", "另外", "此外", "除此之外", "总的来说", "综上所述",
            "第一", "第二", "第三", "第四", "第五", "一方面", "另一方面"
        ]

        # 检查句子是否以段落开始的标志词开头
        sentence_start = sentence.strip()[:10]  # 只检查前10个字符
        for starter in paragraph_starters:
            if starter in sentence_start:
                return True

        # 检查是否有说话人标记（如"张三："或"李四:"）
        if re.search(r'^[A-Za-z\u4e00-\u9fa5]+[：:]', sentence.strip()):
            return True

        # 检查是否是问句开头（通常是新话题）
        if sentence.strip().startswith(("为什么", "怎么", "如何", "什么", "谁", "何时", "何地", "哪里")):
            return True

        return False

    def _merge_short_segments(self, segments):
        """
        合并过短的段落，并分割过长的段落

        参数:
            segments (list): 段落列表

        返回:
            list: 处理后的段落列表
        """
        if not segments:
            return []

        # 理想段落长度范围
        ideal_min_length = 250
        ideal_max_length = 400

        # 第一轮：合并过短的段落
        merged_segments = []
        current_segment = segments[0]

        for i in range(1, len(segments)):
            # 如果当前段落太短，且下一段不是明确的段落开始，则尝试合并
            if len(current_segment) < ideal_min_length and not self._is_paragraph_start(segments[i]):
                current_segment += " " + segments[i]
            # 如果当前段落加上下一个段落的长度仍在理想范围内，且下一段不是明确的段落开始，则合并
            elif len(current_segment) + len(segments[i]) <= ideal_max_length and not self._is_paragraph_start(
                    segments[i]):
                current_segment += " " + segments[i]
            else:
                # 当前段落已经合适或下一段是新段落的开始
                merged_segments.append(current_segment)
                current_segment = segments[i]

        merged_segments.append(current_segment)

        # 第二轮：分割过长的段落
        final_segments = []
        for segment in merged_segments:
            if len(segment) > ideal_max_length + 100:  # 允许稍微超过理想最大长度
                # 按句子分割段落
                sentences = re.split(r'([。！？.!?])', segment)
                current_part = ""

                for i in range(0, len(sentences), 2):
                    if i + 1 < len(sentences):
                        sentence = sentences[i] + sentences[i + 1]
                    else:
                        sentence = sentences[i]

                    # 如果添加这个句子会使段落过长，则开始新段落
                    if len(current_part) + len(sentence) > ideal_max_length and len(current_part) >= ideal_min_length:
                        # 检查是否是段落的自然结束（如句号结尾）
                        if current_part.rstrip().endswith(("。", ".", "!", "?", "！", "？")):
                            final_segments.append(current_part)
                            current_part = sentence
                        else:
                            # 如果不是自然结束，且下一句是段落开始，则分段
                            if self._is_paragraph_start(sentence):
                                final_segments.append(current_part)
                                current_part = sentence
                            else:
                                # 否则继续添加
                                current_part += sentence
                    else:
                        current_part += sentence

                if current_part:
                    final_segments.append(current_part)
            else:
                final_segments.append(segment)

        return final_segments

    def segment_text(self, text):
        """
        对文本进行分段

        参数:
            text (str): 输入文本

        返回:
            list: 分段后的文本列表
        """
        # 如果文本为空，直接返回空列表
        if not text or len(text.strip()) == 0:
            return []

        # 预处理文本
        tokenized_sentences, processed_sentences = self.preprocess_text(text)

        # 如果只有一个句子，直接返回原文本
        if len(processed_sentences) <= 1:
            return [text]

        # 原始句子列表（未分词）
        original_sentences = []
        for paragraph in text.split('\n'):
            if paragraph.strip():
                for sent in re.split(r'(?<=[.。!！?？;；])\s*', paragraph):
                    if sent.strip():
                        original_sentences.append(sent.strip())

        # 确保句子数量一致
        if len(original_sentences) != len(processed_sentences):
            if self.verbose:
                print(
                    f"警告: 原始句子数量 ({len(original_sentences)}) 与处理后句子数量 ({len(processed_sentences)}) 不一致")
            # 使用较短的长度
            min_len = min(len(original_sentences), len(processed_sentences))
            original_sentences = original_sentences[:min_len]
            processed_sentences = processed_sentences[:min_len]

        # 构建LDA主题模型
        vectorizer, lda_model, doc_topic_matrix = self.build_lda_model(processed_sentences)

        # 根据主题变化对文本进行分段
        segments = self._segment_text_by_topic(original_sentences, doc_topic_matrix)

        return segments

    def process_file(self, input_file, output_file=None):
        """
        处理单个文件

        参数:
            input_file (str): 输入文件路径
            output_file (str): 输出文件路径，如果为None，则自动生成

        返回:
            str: 输出文件路径
        """
        if not output_file:
            # 自动生成输出文件路径
            base_name = os.path.basename(input_file)
            name_without_ext = os.path.splitext(base_name)[0]
            output_dir = os.path.dirname(input_file)
            output_file = os.path.join(output_dir, f"{name_without_ext}_segmented.docx")

        # 读取输入文件
        if input_file.lower().endswith('.docx'):
            doc = Document(input_file)
            text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        else:  # 假设是文本文件
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()

        # 分段
        segments = self.segment_text(text)

        # 保存分段结果
        doc = Document()
        for segment in segments:
            doc.add_paragraph(segment)
        doc.save(output_file)

        if self.verbose:
            print(f"处理完成: {input_file} -> {output_file}")
            print(f"原始文本长度: {len(text)} 字符, 分段数量: {len(segments)}")

        return output_file

    def process_folder(self, input_folder, output_folder=None):
        """
        处理文件夹中的所有文件

        参数:
            input_folder (str): 输入文件夹路径
            output_folder (str): 输出文件夹路径，如果为None，则使用输入文件夹

        返回:
            list: 输出文件路径列表
        """
        if not output_folder:
            output_folder = input_folder

        # 确保输出文件夹存在
        os.makedirs(output_folder, exist_ok=True)

        output_files = []

        # 如果输入是文件而非文件夹
        if os.path.isfile(input_folder):
            base_name = os.path.basename(input_folder)
            name_without_ext = os.path.splitext(base_name)[0]
            output_file = os.path.join(output_folder, f"{name_without_ext}_segmented.docx")
            output_files.append(self.process_file(input_folder, output_file))
            return output_files

        # 处理文件夹中的所有文件
        for file_name in os.listdir(input_folder):
            file_path = os.path.join(input_folder, file_name)

            # 只处理docx和txt文件
            if os.path.isfile(file_path) and (
                    file_name.lower().endswith('.docx') or file_name.lower().endswith('.txt')):
                name_without_ext = os.path.splitext(file_name)[0]
                output_file = os.path.join(output_folder, f"{name_without_ext}_segmented.docx")
                output_files.append(self.process_file(file_path, output_file))

        return output_files





if __name__ == "__main__":
    # 指定nltk_data的目录
    nltk_data_path = 'F:/python/nltk_data'
    segmenter = LDASegmenter(nltk_data_path=nltk_data_path)
    input_file = "胰岛素冰箱冷藏paraformer4.docx"
    outfile = segmenter.process_file(input_file)
    print(outfile)
