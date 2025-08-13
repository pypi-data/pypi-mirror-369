import openpyxl
import docx
import shutil
import os
import fiona
import pandas as pd
import re
from enum import Enum
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph as DocxParagraph
from docx import Document as DocxDocument
from spire.doc import Document as SpireDocument, Paragraph as SpireParagraph, OfficeMath, FileFormat


class TextType(Enum):
    Text = 0
    Formula = 1


def fileName(FilePath):
    """
    获取给定路径所指向文件的文件名。
    :param FilePath: 文件路径
    :return:
    """
    return os.path.basename(FilePath)


def fileFatherPath(FilePath):
    """
    获取给定路径所指向文件的父文件夹的名称。
    :param FilePath: 路径
    :return:
    """
    return os.path.dirname(FilePath)


def copyFile(FileOriginalPath, FileNewPath):
    """
    本函数用于复制文件，在复制的同时可以根据需求修改函数名字。通常与for循环结合进行批量复制并改名的操作。

    注意，这里文件路径为包含文件名的路径，可以是相对路径，也可以是绝对路径。如：(1)D:Example/EasyWork/example.png;(2)example.png。

    :param FileOriginalPath: 文件原始路径
    :param FileNewPath: 文件新路径
    :return:
    """
    FileNewFatherPath = os.path.dirname(FileNewPath)
    if not os.path.exists(FileNewFatherPath):
        os.makedirs(FileNewFatherPath)
    shutil.copy2(FileOriginalPath, FileNewPath)


def moveFile(FileOriginalPath, FileNewPath):
    """
    本函数用于剪切文件，在剪切的同时可以根据需求修改文件名字。通常与for循环结合进行批量剪切并改名的操作。

    注意，这里文件路径为包含文件名的路径，可以是相对路径，也可以是绝对路径。如：(1)D:Example/EasyWork/example.png;(2)example.png。

    :param FileOriginalPath: 文件原始路径
    :param FileNewPath: 文件新路径
    :return:
    """
    FileNewFatherPath = os.path.dirname(FileNewPath)
    if not os.path.exists(FileNewFatherPath):
        os.makedirs(FileNewFatherPath)
    shutil.move(FileOriginalPath, FileNewPath)


def excleRead(ExclePath, SheetIndex, Rowlow, Rowmax, Collow, Colmax):
    """
    读取Excle文件(.xlsx/.xls)的函数。

    注意：

    1.这里所有的序号均是从1开始而不是0！而且列号为数字，请不要填写字母。

    2.文件路径同样为包含文件名的路径，可以是相对路径，也可以是绝对路径。

    :param ExclePath: Excle路径
    :param SheetIndex: Sheet序号
    :param Rowlow: 最小行号
    :param Rowmax: 最大行号
    :param Collow: 最小列号
    :param Colmax: 最大列号
    :return: 数据表或数据，数据表是二维列表，数据与单元格数据的性质相同。
    """
    if Rowlow > Rowmax:
        t = Rowmax
        Rowmax = Rowlow
        Rowlow = t
    if Collow > Colmax:
        t = Rowmax
        Colmax = Collow
        Collow = t
    RowNum = Rowmax - Rowlow + 1
    ColNum = Colmax - Collow + 1
    # 打开工作簿
    workbook = openpyxl.load_workbook(ExclePath)
    # 获取所有工作表
    sheets = workbook.sheetnames
    # 选择第一个工作表
    sheet = workbook[sheets[SheetIndex - 1]]
    # 存储为列表

    m = 0
    i = 0
    SheetData = [[None for j in range(Colmax - Collow + 1)] for i in range(Rowmax - Rowlow + 1)]
    for row in sheet.iter_rows():
        n = 0
        j = 0
        for cell in row:

            if Rowlow <= m + 1 <= Rowmax and Collow <= n + 1 <= Colmax:
                # 获取单元格的值
                a = cell.value
                SheetData[i][j] = a
                j = j + 1
            n = n + 1
        if Rowlow <= m + 1 <= Rowmax:
            i = i + 1
        m = m + 1
    if RowNum == 1:
        SheetSingle = list()
        i = 0
        for data in SheetData[0]:
            SheetSingle.append(data)
            i = i + 1
        if ColNum == 1:
            SheetData = SheetSingle[0]
        else:
            SheetData = SheetSingle
    elif ColNum == 1:
        SheetSingle = list()
        for i in range(0, RowNum):
            SheetSingle.append(SheetData[i][0])
        SheetData = SheetSingle
    return SheetData


def excleWrite(ExclePath, SheetIndex, Row, Col, Value, SaveAsNewFile):
    """
    写入Excle文件(.xlsx/.xls)的函数。

    注意：

    1.这里所有的序号均是从1开始而不是0！而且列号为数字，请不要填写字母。

    2.文件路径同样为包含文件名的路径，可以是相对路径，也可以是绝对路径。
    :param ExclePath: Excle路径
    :param SheetIndex: Sheet序号
    :param Row: 单元格行号
    :param Col: 单元格列号
    :param Value: 要赋的值
    :param SaveAsNewFile: 是否保存为新文件(True/False)
    :return:
    """
    if SaveAsNewFile:
        FileNewName = "New_" + fileName(ExclePath)
        if fileFatherPath(ExclePath) != "":
            ExcleNewPath = fileFatherPath(ExclePath) + "\\" + FileNewName
        else:
            ExcleNewPath = FileNewName
        workbook = openpyxl.load_workbook(ExclePath)
        sheet_names = workbook.sheetnames
        SheetName = sheet_names[SheetIndex - 1]
        # 选择要操作的工作表
        sheet = workbook[SheetName]
        # 在指定的单元格写入数据
        sheet.cell(row=Row, column=Col, value=Value)
        # 保存文件
        workbook.save(ExcleNewPath)
    else:
        workbook = openpyxl.load_workbook(ExclePath)
        sheet_names = workbook.sheetnames
        SheetName = sheet_names[SheetIndex - 1]
        # 选择要操作的工作表
        sheet = workbook[SheetName]
        # 在指定的单元格写入数据
        sheet.cell(row=Row, column=Col, value=Value)
        # 保存文件
        workbook.save(ExclePath)


def wordTableRead(WordPath, TableIndex):
    """
    读取Word表格的函数。

    注意：这里表格索引为全局索引。文件路径同样为包含文件名的路径，可以是相对路径，也可以是绝对路径，与前面的函数所需的路径形式相同。
    :param WordPath: Word路径
    :param TableIndex: 表格索引
    :return:
    """
    doc = docx.Document(WordPath)
    table = doc.tables[TableIndex - 1]
    RowNum = 0
    for row in table.rows:
        ColNum = 0
        for cell in row.cells:
            ColNum = ColNum + 1
        RowNum = RowNum + 1
    SheetData = [[None for j in range(ColNum)] for i in range(RowNum)]
    i = 0
    for row in table.rows:
        j = 0
        for cell in row.cells:
            if i == 0 and j == 0:
                cell_text = cell.text
                SheetData[i][j] = cell_text
            else:
                bcell_text = cell_text
                cell_text = cell.text
                if bcell_text != cell_text:
                    SheetData[i][j] = cell_text
                else:
                    SheetData[i][j] = None
            j = j + 1
        i = i + 1
    return SheetData


def wordTableWrite(WordPath, TableIndex, Row, Col, Text, SaveAsNewFile):
    """
    写入Word表格的函数。

    注意：这里行号与Excle的不同，加入表格1的未合并前为6个单元格，此时将1、2单元格合并。此时“行号”参数填写1与2均会写入第一个单元格，当填入3时才会写入第二个单元格。列与行的情况相同。

    如果想要插入如下的上标下标：

    from FreeWork import OfficeWork as ow

    ow.WordTableWrite(WordPath, TableIndex, Row, Col, "面积 S_(1)=123 hm^(2)", SaveAsNewFile(True / False))

    其中括号是必不可少的，否则“^”符号后面的所有文本均将以上标的形式写入段落，“_”符号后面的所有文本均将以下标的形式写入段落，直至本条插入文本结束！还有请注意，这里括号需以英文状态下输入，否则将不会起到其应有的作用。
    :param WordPath: Word路径
    :param TableIndex: 表格索引
    :param Row: 行号
    :param Col: 列号
    :param Text: 欲写入的文本
    :param SaveAsNewFile: 是否保存为新文件(True/False)
    :return:
    """
    if SaveAsNewFile:
        FileNewName = "New_" + fileName(WordPath)
        if fileFatherPath(WordPath) != "":
            WordNewPath = fileFatherPath(WordPath) + "\\" + FileNewName
        else:
            WordNewPath = FileNewName
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        ## 新增代码段 ##
        SuperTextBefor = ""
        SuperText = ""
        SuperTextAfter = ""
        SubTextBefor = ""
        SubText = ""
        SubTextAfter = ""
        for i in range(0, len(Text)):
            if Text[i] == "^" and i + 1 < len(Text):
                SubTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SuperTextAfter = SuperTextAfter + Text[k]
                            break
                        SuperText = SuperText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SuperText = SuperText + Text[j]
                    break
            elif Text[i] == "_" and i + 1 < len(Text):
                SuperTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SubTextAfter = SubTextAfter + Text[k]
                            break
                        SubText = SubText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SubText = SubText + Text[j]
                    break
            SuperTextBefor = SuperTextBefor + Text[i]
            SubTextBefor = SubTextBefor + Text[i]

        if SuperTextBefor == Text and SubTextBefor == Text:
            Cell.text = Text
        elif SuperText != "":
            Cell.text = SuperTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SuperText)
            AddedSuperText.font.superscript = True
            paragraph.add_run(SuperTextAfter)
        else:
            Cell.text = SubTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SubText)
            AddedSuperText.font.subscript = True
            paragraph.add_run(SubTextAfter)
        ## 新增代码段 ##
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordNewPath)
    else:
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        ## 新增代码段 ##
        SuperTextBefor = ""
        SuperText = ""
        SuperTextAfter = ""
        SubTextBefor = ""
        SubText = ""
        SubTextAfter = ""
        for i in range(0, len(Text)):
            if Text[i] == "^" and i + 1 < len(Text):
                SubTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SuperTextAfter = SuperTextAfter + Text[k]
                            break
                        SuperText = SuperText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SuperText = SuperText + Text[j]
                    break
            elif Text[i] == "_" and i + 1 < len(Text):
                SuperTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SubTextAfter = SubTextAfter + Text[k]
                            break
                        SubText = SubText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SubText = SubText + Text[j]
                    break
            SuperTextBefor = SuperTextBefor + Text[i]
            SubTextBefor = SubTextBefor + Text[i]

        if SuperTextBefor == Text and SubTextBefor == Text:
            Cell.text = Text
        elif SuperText != "":
            Cell.text = SuperTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SuperText)
            AddedSuperText.font.superscript = True
            paragraph.add_run(SuperTextAfter)
        else:
            Cell.text = SubTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SubText)
            AddedSuperText.font.subscript = True
            paragraph.add_run(SubTextAfter)
        ## 新增代码段 ##
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordPath)


def wordTableInsertFig(WordPath, TableIndex, Row, Col, ImagePath, ImageHeight_cm, ImageWidth_cm, SaveAsNewFile):
    """
    Word表格追加图片函数，不删除原有文字。

    注意：这里图片高度可以为“None”。
    :param WordPath: Word路径
    :param TableIndex: 表格索引
    :param Row: 行号
    :param Col: 列号
    :param ImagePath: 图片路径
    :param ImageHeight_cm: 插入后图片的高度（厘米为单位）
    :param ImageWidth_cm: 插入后图片的宽度（厘米为单位）
    :param SaveAsNewFile: 是否保存为新文件(True/False)
    :return:
    """
    if SaveAsNewFile:
        FileNewName = "New_" + fileName(WordPath)
        if fileFatherPath(WordPath) != "":
            WordNewPath = fileFatherPath(WordPath) + "\\" + FileNewName
        else:
            WordNewPath = FileNewName
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        if ImageHeight_cm is not None and ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm), width=Cm(ImageWidth_cm))
        elif ImageHeight_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm))
        elif ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, width=Cm(ImageWidth_cm))
        else:
            Cell.add_paragraph().add_run().add_picture(ImagePath)
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordNewPath)
    else:
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        if ImageHeight_cm is not None and ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm), width=Cm(ImageWidth_cm))
        elif ImageHeight_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm))
        elif ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, width=Cm(ImageWidth_cm))
        else:
            Cell.add_paragraph().add_run().add_picture(ImagePath)
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordPath)


def wordTableParaAlignment(WordPath, TableIndex, Row, Col, Alignment_left_right_center_None, SaveAsNewFile):
    """
    Word表格单元格对齐设置函数。

    注意：对齐方式只能填写left/right/center/None，否则均会设置为None两端对齐。
    :param WordPath: Word路径
    :param TableIndex: 表格索引
    :param Row: 行号
    :param Col: 列号
    :param Alignment_left_right_center_None: 对齐方式
    :param SaveAsNewFile: 是否保存为新文件(True/False)
    :return:
    """
    if SaveAsNewFile:
        FileNewName = "New_" + fileName(WordPath)
        if fileFatherPath(WordPath) != "":
            WordNewPath = fileFatherPath(WordPath) + "\\" + FileNewName
        else:
            WordNewPath = FileNewName
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        if Alignment_left_right_center_None == "left":
            alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif Alignment_left_right_center_None == "right":
            alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif Alignment_left_right_center_None == "center":
            alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            alignment = None
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignment  # 设置为原段落的对齐方式
        doc.save(WordNewPath)
    else:
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        if Alignment_left_right_center_None == "left":
            alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif Alignment_left_right_center_None == "right":
            alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif Alignment_left_right_center_None == "center":
            alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            alignment = None
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignment  # 设置为原段落的对齐方式
        doc.save(WordPath)


def ShpToXlsx(ShpPath, XlsxPath):
    """
    Shapefile转出Excle函数。
    :param ShpPath: Shp路径
    :param XlsxPath: Xlsx路径
    :return:
    """
    with fiona.open(ShpPath, "r") as shapefile:
        # 读取SHP文件的属性表
        properties = [feature["properties"] for feature in shapefile]

        # 将属性表转换为DataFrame
        df = pd.DataFrame(properties)

        # 将DataFrame写入Excel文件
        df.to_excel(XlsxPath, index=False)


def wordAdd(wordPath, wordSavePath, new_text, FontName=None, FontSize=None, IsBold=None, IsItalic=None):
    """
    在文档的末尾添加文字(不开新段落)。
    :param wordPath: Word路径
    :param wordSavePath: Word保存路径
    :param new_text: 添加文本
    :param FontName: 字体名称
    :param FontSize: 字体大小
    :param IsBold: 是否加粗
    :param IsItalic: 是否斜体
    :return:
    """
    doc = docx.Document(wordPath)
    # 获取文档的最后一段
    last_paragraph = doc.paragraphs[-1]

    if last_paragraph.runs:
        # 获取最后一段的最后一个run
        last_run = last_paragraph.runs[-1]
        # 创建一个新的run
        new_run = last_paragraph.add_run(new_text)

        # 复制字体名称
        if FontName is None:
            if last_run.font.name:
                new_run.font.name = last_run.font.name
                new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), last_run.font.name)
        else:
            new_run.font.name = FontName
            new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), FontName)
        # 复制字体大小
        if FontSize is None:
            if last_run.font.size:
                new_run.font.size = last_run.font.size
        else:
            new_run.font.size = FontSize
        # 复制加粗、斜体等其他属性
        if IsBold is None:
            new_run.bold = last_run.bold
        else:
            if IsBold:
                new_run.bold = True
            else:
                new_run.bold = False
        if IsItalic is None:
            new_run.italic = last_run.italic
        else:
            if IsItalic:
                new_run.italic = True
            else:
                new_run.italic = False
    else:
        # 如果最后一段是空的，则直接添加新文本
        last_paragraph.add_run(new_text)

    # 保存修改后的文档
    doc.save(wordSavePath)


def wordParagraphAdd(wordPath, wordSavePath, new_text, FontName=None, FontSize=None, IsBold=None, IsItalic=None,
                     Indent=None, Alignment="l"):
    """
    在文档的末尾添加文字(开新段落)。
    :param wordPath: Word路径
    :param wordSavePath: Word保存路径
    :param new_text: 添加文本
    :param FontName: 字体名称
    :param FontSize: 字体大小
    :param IsBold: 是否加粗
    :param IsItalic: 是否斜体
    :param Indent: 缩进值
    :param Alignment: 对其方式(l/c/r/j)
    :return:
    """
    doc = docx.Document(wordPath)
    # 获取文档的最后一段
    last_paragraph = doc.paragraphs[-1]

    if last_paragraph.runs:
        # 获取最后一段的最后一个run
        last_run = last_paragraph.runs[-1]
        # 创建一个新的run
        new_paragraph = doc.add_paragraph()
        new_run = new_paragraph.add_run(new_text)

        format = last_paragraph.paragraph_format

        # 获取左缩进和首行缩进（单位是英寸，Word中的默认单位）
        left_indent = format.left_indent
        first_line_indent = format.first_line_indent
        if Indent is None:
            new_paragraph.paragraph_format.first_line_indent = first_line_indent
        else:
            try:
                new_paragraph.paragraph_format.first_line_indent = Indent
            except:
                new_paragraph.paragraph_format.first_line_indent = first_line_indent
        if Alignment == "l":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif Alignment == "c":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif Alignment == "r":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif Alignment == "j":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 复制字体名称
        if FontName is None:
            if last_run.font.name:
                new_run.font.name = last_run.font.name
                new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), last_run.font.name)
        else:
            new_run.font.name = FontName
            new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), FontName)
        # 复制字体大小
        if FontSize is None:
            if last_run.font.size:
                new_run.font.size = last_run.font.size
        else:
            new_run.font.size = FontSize
        # 复制加粗、斜体等其他属性
        if IsBold is None:
            new_run.bold = last_run.bold
        else:
            if IsBold:
                new_run.bold = True
            else:
                new_run.bold = False
        if IsItalic is None:
            new_run.italic = last_run.italic
        else:
            if IsItalic:
                new_run.italic = True
            else:
                new_run.italic = False
    else:
        # 如果最后一段是空的，则直接添加新文本
        last_paragraph.add_run(new_text)

    # 保存修改后的文档
    doc.save(wordSavePath)


def wordParaFormat(wordPath):
    doc = docx.Document(wordPath)
    # 获取文档的最后一段
    last_paragraph = doc.paragraphs[-1]
    last_run = last_paragraph.runs[-1]
    format = last_paragraph.paragraph_format
    return format


def insert_paragraph_after(paragraph: DocxParagraph, text=None, style=None) -> DocxParagraph:
    """
    在指定 python-docx 段落之后插入新段落。

    :param paragraph: DocxParagraph
        目标段落对象，表示在该段落后插入新段落。
    :param text: str or None
        新段落的文本内容，如果为 None，则新段落为空。
    :param style: str or None
        新段落样式名称，须是 Word 中已存在的样式名称，例如“标题 1”。
        若为 None，则使用默认样式。

    :return: DocxParagraph
        新插入的段落对象。
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = DocxParagraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def wordInsertText(WordPath, Text, ParaIndex, NewParagraph=True, StyleName=None):
    """
    使用 python-docx 库向指定的 .docx 文件中指定段落插入文本。

    功能：
    - 支持在指定段落索引之后新建段落插入文本，或在指定段落末尾追加文本；
    - 若段落索引越界则自动调整为可用范围的索引；
    - 可设置插入文本段落样式；
    - 若文档为空，则自动创建第一段落插入文本；
    - 操作完成后覆盖保存原文件。

    :param WordPath: str
        Word文档路径（.docx格式），路径需存在。
    :param Text: str
        要插入的文本内容。
    :param ParaIndex: int
        目标段落索引（0开始），文本将插入该段落之后（新段落）或该段落末尾（追加文本）。
    :param NewParagraph: bool, 默认 True
        是否新建段落插入文本：
        - True：目标段落后新建段落写入文本；
        - False：不新建段落，追加到目标段落末尾。
    :param StyleName: str or None
        段落样式名称，须在Word文档中已定义，若为 None 使用默认样式。

    :return: None
        完成后覆盖保存 Word 文件。
    """
    doc = DocxDocument(WordPath)
    para_count = len(doc.paragraphs)

    if para_count == 0:
        # 文档无段落，直接添加新段落
        new_para = doc.add_paragraph(Text)
        if StyleName:
            new_para.style = StyleName
        doc.save(WordPath)
        return

    if ParaIndex < 0:
        ParaIndex = 0
    elif ParaIndex >= para_count:
        ParaIndex = para_count - 1

    target_para = doc.paragraphs[ParaIndex]

    if NewParagraph:
        insert_paragraph_after(target_para, text=Text, style=StyleName)
    else:
        run = target_para.add_run(Text)
        if StyleName:
            target_para.style = StyleName

    doc.save(WordPath)


def wordInsertLatexFormula(WordPath, ParaIndex, LatexCode, NewParagraph=True):
    """
    使用 Spire.Doc 库向 .docx 文件指定段落插入 LaTeX 数学公式。

    功能：
    - 加载指定路径 Word 文档（.docx格式）；
    - 根据参数决定是在指定段落后新建段落插入公式，还是追加到指定段落末尾；
    - 若文档无段落，自动创建第一个空段落后插入公式；
    - 在插入公式时调用 Spire.Doc OfficeMath 支持 LaTeX；
    - 保存时覆盖保存原文件。

    :param WordPath: str
        Word文档路径（需为有效的 .docx 文件），操作后覆盖保存。
    :param ParaIndex: int
        目标段落索引（0开始），插入公式位置依参数决定。
    :param LatexCode: str
        LaTeX 数学公式代码（标准 LaTeX 格式）。
    :param NewParagraph: bool, 默认 True
        是否新建段落：
        - True：新建段落插入公式；
        - False：追加到指定段落末尾。

    :return: None
        完成后保存 Word 文档。
    """
    doc = SpireDocument()
    doc.LoadFromFile(WordPath)

    section = doc.Sections[0]
    para_count = section.Paragraphs.Count

    if para_count == 0:
        new_para = SpireParagraph(doc)
        section.Paragraphs.Add(new_para)
        para_count = 1

    if ParaIndex < 0:
        ParaIndex = 0
    elif ParaIndex >= para_count:
        ParaIndex = para_count - 1

    target_para = section.Paragraphs[ParaIndex]

    if NewParagraph:
        new_para = SpireParagraph(doc)
        section.Paragraphs.Insert(ParaIndex + 1, new_para)
        math_obj = OfficeMath(doc)
        math_obj.FromLatexMathCode(LatexCode)
        new_para.Items.Add(math_obj)
    else:
        math_obj = OfficeMath(doc)
        math_obj.FromLatexMathCode(LatexCode)
        target_para.Items.Add(math_obj)

    doc.SaveToFile(WordPath, FileFormat.Docx2016)
    doc.Dispose()


def wordTableWriteExtend(WordPath, TableIndex, Row, Col, Text, StyleName=None,
                         Alignment: WD_ALIGN_PARAGRAPH | None = None):
    """
    向Word指定表格的位置写入文本，可自动新增行（新增行复制格式含对齐但不复制文字）
    支持段落样式和写入后目标对齐设置。

    :param WordPath: Word文档路径(.docx)
    :param TableIndex: 表格序号，从1开始
    :param Row: 行号，从1开始
    :param Col: 列号，从1开始
    :param Text: 写入文本
    :param StyleName: 段落样式名（字符串，可选）
    :param Alignment: 对齐方式(WD_ALIGN_PARAGRAPH枚举），None则保持原对齐
    """
    doc = docx.Document(WordPath)
    tables = doc.tables

    # 校验表格索引
    if TableIndex < 1 or TableIndex > len(tables):
        raise IndexError(f"表格索引 TableIndex={TableIndex} 超出范围，有效值：1~{len(tables)}")
    table = tables[TableIndex - 1]

    # 校验列索引
    if Col < 1 or Col > len(table.columns):
        raise IndexError(f"列索引 Col={Col} 超出范围，表格列数：{len(table.columns)}")

    # 新增空行，复制上一行格式(含段落对齐)
    if Row > len(table.rows):
        last_row = table.rows[-1]
        for _ in range(Row - len(table.rows)):
            new_row = table.add_row()

            # 先清空新增行所有单元格文字和runs
            for cell in new_row.cells:
                cell.text = ''
                for para in cell.paragraphs:
                    while para.runs:
                        para._element.remove(para.runs[0]._element)

            # 复制格式但不复制文字
            for idx, src_cell in enumerate(last_row.cells):
                dest_cell = new_row.cells[idx]
                for para_idx, src_para in enumerate(src_cell.paragraphs):
                    # 确保目标段落足够
                    while para_idx >= len(dest_cell.paragraphs):
                        dest_cell.add_paragraph()

                    dest_para = dest_cell.paragraphs[para_idx]
                    dest_para.clear()

                    # 复制字体格式，不复制文字
                    for src_run in src_para.runs:
                        new_run = dest_para.add_run('')
                        if src_run.font.name:
                            new_run.font.name = src_run.font.name
                            new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), src_run.font.name)
                        if src_run.font.size:
                            new_run.font.size = src_run.font.size
                        new_run.bold = src_run.bold
                        new_run.italic = src_run.italic
                        new_run.underline = src_run.underline

                    # 复制段落格式（包含对齐）
                    pf = src_para.paragraph_format
                    dest_para.paragraph_format.left_indent = pf.left_indent
                    dest_para.paragraph_format.first_line_indent = pf.first_line_indent
                    dest_para.paragraph_format.space_before = pf.space_before
                    dest_para.paragraph_format.space_after = pf.space_after
                    dest_para.paragraph_format.alignment = pf.alignment

    # 定位目标单元格
    cell = table.cell(Row - 1, Col - 1)

    # 获取第一个段落或新建
    if cell.paragraphs:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()

    # 读取原对齐
    ori_align = para.paragraph_format.alignment

    # 清空段落所有runs保留段落元素
    while para.runs:
        para._element.remove(para.runs[0]._element)

    # 写入文本run
    run = para.add_run(Text)

    # 设置段落样式（可选）
    if StyleName:
        para.style = StyleName

    # 设置对齐，传入优先，否则恢复原对齐
    if Alignment is not None:
        para.paragraph_format.alignment = Alignment
    else:
        para.paragraph_format.alignment = ori_align

    # 最后保存文档
    doc.save(WordPath)


def wordTableInsertLatexFormula(WordPath, TableIndex, Row, Col, LatexCode, NewParagraph=False):
    """
    使用 Spire.Doc 库向 .docx 文件指定表格单元格插入 LaTeX 数学公式。

    功能：
    - 加载指定路径 Word 文档（.docx格式）；
    - 根据参数决定是在指定单元格后新建段落插入公式，还是追加到指定单元格内；
    - 若单元格无段落，自动创建第一个空段落后插入公式；
    - 在插入公式时调用 Spire.Doc OfficeMath 支持 LaTeX；
    - 保存时覆盖保存原文件。

    :param WordPath: str
        Word文档路径（需为有效的 .docx 文件），操作后覆盖保存。
    :param TableIndex: int
        表格索引，从1开始。
    :param Row: int
        行号，从1开始。
    :param Col: int
        列号，从1开始。
    :param LatexCode: str
        LaTeX 数学公式代码（标准 LaTeX 格式）。
    :param NewParagraph: bool, 默认 True
        是否新建段落：
        - True：新建段落插入公式；
        - False：追加到指定单元格末尾现有段落。

    :return: None
        完成后保存 Word 文档。
    """
    # 使用 Spire.Doc 打开文档
    doc = SpireDocument()
    doc.LoadFromFile(WordPath)

    # 获取指定的表格
    table = doc.Sections[0].Tables[TableIndex - 1]

    # 获取目标单元格
    cell = table.Rows[Row - 1].Cells[Col - 1]

    # 获取单元格中的段落，如果没有段落则创建一个
    if cell.Paragraphs.Count == 0:
        paragraph = cell.AddParagraph()
    else:
        paragraph = cell.Paragraphs[cell.Paragraphs.Count - 1]  # 获取最后一个段落

    # 如果需要新建段落插入公式
    if NewParagraph:
        new_para = SpireParagraph(doc)
        cell.Paragraphs.Add(new_para)
        math_obj = OfficeMath(doc)
        math_obj.FromLatexMathCode(LatexCode)
        new_para.Items.Add(math_obj)
    else:
        # 追加公式到现有段落
        math_obj = OfficeMath(doc)
        math_obj.FromLatexMathCode(LatexCode)
        paragraph.Items.Add(math_obj)

    # 保存文档
    doc.SaveToFile(WordPath, FileFormat.Docx2016)

    # 释放资源
    doc.Dispose()


def wordTableInsertText(WordPath, TableIndex, Row, Col, Text, StyleName=None, NewParagraph=False):
    """
    向Word表格的指定单元格插入普通文本，并保持原有单元格格式不变。
    :param WordPath: Word文档路径（.docx格式）
    :param TableIndex: 表格索引，从1开始
    :param Row: 行号，从1开始
    :param Col: 列号，从1开始
    :param Text: 要插入的文本
    :param StyleName: 段落样式名称，可选
    :param NewParagraph: bool, 默认 False
        是否新建段落插入文本：
        - True：新建段落插入文本；
        - False：追加到指定单元格的现有段落。
    """
    doc = docx.Document(WordPath)
    table = doc.tables[TableIndex - 1]

    # 获取指定单元格
    cell = table.cell(Row - 1, Col - 1)

    # 如果NewParagraph为True，创建新段落，否则获取最后一个段落
    if NewParagraph:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()

    # 保存原始段落的格式（如字体、大小等）
    original_alignment = paragraph.paragraph_format.alignment
    original_left_indent = paragraph.paragraph_format.left_indent
    original_first_line_indent = paragraph.paragraph_format.first_line_indent
    original_space_before = paragraph.paragraph_format.space_before
    original_space_after = paragraph.paragraph_format.space_after

    # 添加文本
    run = paragraph.add_run(Text)

    # 恢复原始段落格式
    paragraph.paragraph_format.alignment = original_alignment
    paragraph.paragraph_format.left_indent = original_left_indent
    paragraph.paragraph_format.first_line_indent = original_first_line_indent
    paragraph.paragraph_format.space_before = original_space_before
    paragraph.paragraph_format.space_after = original_space_after

    # 设置段落样式（如果提供了样式）
    if StyleName:
        paragraph.style = StyleName

    # 覆盖保存原文件
    doc.save(WordPath)


def wordFindTableIndexByMark(WordPath, MarkRow, MarkString):
    """
    查找Word文档中所有在指定行包含指定标记字符串的表格索引（从1开始）。

    :param WordPath: Word文件路径（.docx）
    :param MarkRow: 标记行号（从1开始），表示在该行查找标记字符串
    :param MarkString: 要匹配的标记字符串
    :return: List[int]，所有符合条件的表格索引列表，索引从1开始
    """
    doc = docx.Document(WordPath)
    tables = doc.tables
    matched_indexes = []

    for idx, table in enumerate(tables, start=1):
        if MarkRow < 1 or MarkRow > len(table.rows):
            # 如果表格行数不足，跳过
            continue
        row = table.rows[MarkRow - 1]
        # 遍历该行所有单元格
        for cell in row.cells:
            if MarkString in cell.text:
                matched_indexes.append(idx)
                break  # 该表符合条件，跳出单元格循环，继续下一个表格

    return matched_indexes


def wordFindCellsByMark(WordPath, MarkString):
    """
    查找Word文档中所有包含指定标记字符串的单元格，返回列表元素为[表格索引, 行号, 列号]，均从1开始。

    :param WordPath: Word文件路径（.docx）
    :param MarkString: 要匹配的标记字符串
    :return: List[List[int, int, int]]
        其中元素为 [表格索引，从1开始，行号从1开始，列号从1开始]
    """
    doc = docx.Document(WordPath)
    tables = doc.tables
    matched_cells = []

    for table_idx, table in enumerate(tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            for col_idx, cell in enumerate(row.cells, start=1):
                if MarkString in cell.text:
                    matched_cells.append([table_idx, row_idx, col_idx])

    return matched_cells


def wordFindParagraphIndexesByMark(WordPath, MarkString):
    """
    查找Word文档中所有包含指定标记字符串的段落索引列表（从0开始）。

    :param WordPath: Word文件路径（.docx）
    :param MarkString: 要匹配的标记字符串
    :return: List[int]，所有匹配段落的索引列表（从0开始）
    """
    doc = docx.Document(WordPath)
    matched_indexes = []

    for idx, paragraph in enumerate(doc.paragraphs):
        if MarkString in paragraph.text:
            matched_indexes.append(idx)

    return matched_indexes


def wordReplaceParagraphText(WordPath, Text, ParaIndex, StyleName=None):
    """
    使用 python-docx 库替换指定Word文档中指定段落的全部文本。

    功能：
    - 按段落索引替换该段落所有文本为新文本；
    - 段落索引越界时自动调整为有效范围；
    - 支持设置段落样式（段落全部替换）；
    - 若文档无段落，则新建第一段并写入文本；
    - 修改后覆盖保存原文件。

    :param WordPath: str
        Word文档路径（.docx格式），路径需存在。
    :param Text: str
        替换成的新文本内容。
    :param ParaIndex: int
        目标段落索引（0开始）。
    :param StyleName: str or None
        段落样式名称，须在Word文档中已定义，若为 None 使用默认样式。

    :return: None
        完成后覆盖保存 Word 文件。
    """
    doc = DocxDocument(WordPath)
    para_count = len(doc.paragraphs)

    if para_count == 0:
        # 若无段落，添加新段
        new_para = doc.add_paragraph(Text)
        if StyleName:
            new_para.style = StyleName
        doc.save(WordPath)
        return

    # 索引越界修正
    if ParaIndex < 0:
        ParaIndex = 0
    elif ParaIndex >= para_count:
        ParaIndex = para_count - 1

    target_para = doc.paragraphs[ParaIndex]

    # 清空所有runs（全部文本）
    while target_para.runs:
        target_para._element.remove(target_para.runs[0]._element)

    # 添加新文本run
    new_run = target_para.add_run(Text)

    # 设置样式
    if StyleName:
        target_para.style = StyleName

    doc.save(WordPath)


def wordFindParagraphTextsByMark(WordPath, MarkString):
    """
    查找Word文档中所有包含指定标记字符串的段落文本列表。

    :param WordPath: Word文档路径（.docx）
    :param MarkString: 要匹配的标记字符串
    :return: List[str]，所有匹配段落的完整文本列表
    """
    doc = docx.Document(WordPath)
    matched_texts = []

    for paragraph in doc.paragraphs:
        if MarkString in paragraph.text:
            matched_texts.append(paragraph.text)

    return matched_texts


def wordDeleteParagraph(WordPath, ParaIndex):
    """
    删除指定索引的段落，并覆盖保存Word文件。

    :param WordPath: Word文档路径（.docx格式）
    :param ParaIndex: 目标段落索引（0开始）
    :return: None
    """
    doc = docx.Document(WordPath)
    para_count = len(doc.paragraphs)

    if para_count == 0:
        # 文档无段落，无需操作
        return

    # 修正索引范围
    if ParaIndex < 0:
        ParaIndex = 0
    elif ParaIndex >= para_count:
        ParaIndex = para_count - 1

    # 通过段落XML元素移除段落
    p = doc.paragraphs[ParaIndex]._element
    p.getparent().remove(p)
    p._p = p._element = None

    doc.save(WordPath)


def wordDeleteTable(WordPath, TableIndex):
    """
    删除指定索引的表格，并覆盖保存Word文件。

    :param WordPath: Word文档路径（.docx格式）
    :param TableIndex: 表格索引（从1开始）
    :return: None
    """
    doc = docx.Document(WordPath)
    tables = doc.tables
    table_count = len(tables)

    if table_count == 0:
        # 文档无表格，无需操作
        return

    # 修正索引范围
    if TableIndex < 1:
        TableIndex = 1
    elif TableIndex > table_count:
        TableIndex = table_count

    tbl = tables[TableIndex - 1]._element
    tbl.getparent().remove(tbl)
    tbl._tbl = None

    doc.save(WordPath)


def wordDeleteTableRow(WordPath, TableIndex, RowIndex):
    """
    删除Word文档中指定表格的指定行，并覆盖保存文档。

    :param WordPath: Word文档路径（.docx格式）
    :param TableIndex: 表格索引（从1开始）
    :param RowIndex: 行索引（从1开始）
    :return: None
    """
    doc = docx.Document(WordPath)
    tables = doc.tables
    table_count = len(tables)

    if table_count == 0:
        # 文档无表格，无需操作
        return

    # 修正表格索引范围
    if TableIndex < 1:
        TableIndex = 1
    elif TableIndex > table_count:
        TableIndex = table_count

    table = tables[TableIndex - 1]
    rows_num = len(table.rows)

    if rows_num == 0:
        # 表格无行，无需操作
        return

    # 修正行索引范围
    if RowIndex < 1:
        RowIndex = 1
    elif RowIndex > rows_num:
        RowIndex = rows_num

    row_to_delete = RowIndex - 1  # 0-based索引

    row_element = table.rows[row_to_delete]._tr
    parent_element = row_element.getparent()
    parent_element.remove(row_element)

    doc.save(WordPath)


def parse_text_and_formula(text):
    """
    自动解析文本，提取LaTeX公式与普通文本。
    支持以下公式标记：
      - 双美元 $$...$$
      - 单美元 $...$
      - \( ... \)
      - \[ ... \]

    返回列表，每个元素格式为 [类型, 内容]，类型为"text"或"formula"

    :param text: 原始字符串
    :return: list [["text", str], ["formula", str], ...]
    """
    # 匹配LaTeX公式的正则（捕获组）
    pattern = re.compile(
        r'(\$\$.*?\$\$|'  # 双美元 $$...$$
        r'\$.*?\$|'  # 单美元 $...$
        r'\\\(.*?\\\)|'  # \( ... \)
        r'\\\[.*?\\\])',  # \[ ... \]
        re.DOTALL
    )

    result = []
    last_end = 0

    for m in pattern.finditer(text):
        start, end = m.span()

        # 普通文本（匹配到的公式前面）
        if start > last_end:
            normal_text = text[last_end:start]
            if normal_text:
                result.append([TextType.Text, normal_text])

        formula_text = m.group()
        # 去除包裹符号，只保留公式内容
        if formula_text.startswith('$$') and formula_text.endswith('$$'):
            content = formula_text[2:-2].strip()
        elif formula_text.startswith('$') and formula_text.endswith('$'):
            content = formula_text[1:-1].strip()
        elif formula_text.startswith('\\(') and formula_text.endswith('\\)'):
            content = formula_text[2:-2].strip()
        elif formula_text.startswith('\\[') and formula_text.endswith('\\]'):
            content = formula_text[2:-2].strip()
        else:
            content = formula_text

        if content:
            result.append([TextType.Formula, content])

        last_end = end

    # 结尾可能还有普通文本没有匹配公式
    if last_end < len(text):
        tail_text = text[last_end:]
        if tail_text:
            result.append([TextType.Text, tail_text])

    return result
