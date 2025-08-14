# This file is original taken from https://github.com/haesleinhuepf/docx2markdown
# - src/docx2markdown/_docx_to_markdown.py
# It is being altered here to provide results needed for this project.
# original license: https://github.com/haesleinhuepf/docx2markdown/blob/main/LICENSE
# included here for completeness
#---
# BSD 3-Clause License
# Copyright (c) 2024, Robert Haase, ScaDS.AI, Uni Leipzig
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
# 1. Redistributions of source code must retain the above copyright notice, this
#    list of conditions and the following disclaimer.
# 2. Redistributions in binary form must reproduce the above copyright notice,
#    this list of conditions and the following disclaimer in the documentation
#    and/or other materials provided with the distribution.
# 3. Neither the name of the copyright holder nor the names of its
#    contributors may be used to endorse or promote products derived from
#    this software without specific prior written permission.
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#---
# All changes in the this version are Copyright (c) 2025, Paul Philion, Acme Rocket Company
# under the provided MIT license.

import re
import docx
import base64
import os
from lxml import etree
from pathlib import Path, PurePath
import logging

from .page import Page
from .converter import Converter


log = logging.getLogger(__name__)


### new convert file -> memory
def convert(docx_file:Path) -> Page:
    doc = docx.Document(docx_file)

    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)

    markdown = []

    for block in doc.element.body:
        if block.tag.endswith('p'):  # Handle paragraphs
            paragraph = paragraphs.pop(0)  # Match current paragraph
            md_paragraph = ""

            ### switching on paragraph.style.name
            style_name = paragraph.style.name
            #print("STYLE:", style_name)

            if "Heading 1" in style_name:
                md_paragraph = "# "
            elif "Heading 2" in style_name:
                md_paragraph = "## "
            elif "Heading 3" in style_name:
                md_paragraph = "### "
            elif "Heading 4" in style_name:
                md_paragraph = "#### "
            elif "Heading 5" in style_name:
                md_paragraph = "##### "
            elif "Normal" or "normal" in style_name:
                md_paragraph = ""
                if is_list(paragraph):
                    md_paragraph = get_bullet_point_prefix(paragraph)
            else:
                log.error("Unsupported style:", style_name)

            content = parse_run(paragraph)

            md_paragraph += content
            markdown.append(md_paragraph)

        elif block.tag.endswith('tbl'):  # Handle tables (if present)
            table = tables.pop(0)  # Match current table
            table_text = ""
            for i, row in enumerate(table.rows):
                table_text += "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |\n"
                if i == 0:
                    table_text += "| " + " | ".join("---" for _ in row.cells) + " |\n"

            markdown.append(table_text)

        elif block.tag.endswith('sectPr') or block.tag.endswith('sdt'):
            # ignore
            log.debug(f"!!! section ptr: {block}")
            pass
        else:
            log.warning("Unsupported block:", docx_file, block.tag)

    # append footnotes
    markdown.extend(comments(doc))

    # append images to the array of markdown paragraphs
    markdown.extend(embedded_images(doc))

    return Page(
        id = "",
        title = doc.core_properties.title, # docx file metadata
        path = "",
        content = "\n\n".join(markdown),
        editor = "markdown",
        locale = "en",
        tags = doc.core_properties.keywords, # docx file metadata
        description = f"generated from: {docx_file}",
        isPublished = False,
        isPrivate = True,
    )


def comments(doc:docx.Document) -> list[str]:
    comments = []
    for comment in doc.comments:
        # author, text, timestamp
        datestr = comment.timestamp.strftime('%y-%m-%d %H:%M')
        comments.append(f"\n[^{comment.comment_id}]: At {datestr}, {comment.author} said: {comment.text.strip()}")
        log.debug(comment)
    return comments


def write_images(doc:docx.Document, outroot:Path):
    # save all images
    image_folder = Path(outroot, "images")
    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_filename = save_image(rel.target_part, image_folder)
            images[rel.rId] = image_filename[len(outroot)+1:]
            # use relative
            log.info("image file:", rel.rId, ":", image_filename, "-->", images[rel.rId])
        else:
            log.info(f"rel.rId={rel.rId}, rel.reltype={rel.reltype}")
    return images


def embedded_images(doc:docx.Document) -> list[str]:
    """
    Append DOCX images inline in the markdown
    """
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            # rId is in the form "rId18"
            anchor = f"image{rel.rId[3:]}"
            content = base64.b64encode(rel.target_part.blob).decode('utf-8')
            log.info(f"appening image id: {anchor}, size: {len(content)}")
            images.append(f"[{anchor}]: <data:image/png;base64,{content}>\n\n")
        else:
            log.info(f"rel.rId={rel.rId}, rel.reltype={rel.reltype}")

    return images


def extract_r_embed(xml_string):
    """
    Extract the value of r:embed from the given XML string.

    :param xml_string: The XML content as a string.
    :return: The value of r:embed or None if not found.
    """
    # Parse the XML
    root = etree.fromstring(xml_string)

    # Define the namespaces
    namespaces = {
        'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
        'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        'pic': "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }

    # Use XPath to find the <a:blip> element with r:embed
    blip = root.find(".//a:blip", namespaces=namespaces)

    # Extract the r:embed attribute value
    if blip is not None:
        return blip.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    return None


def extract_comment_id(xml_string) -> int:
    """
    Extract the value of w:commentReference w:id="3" the given XML string.
    :param xml_string: The XML content as a string.
    """
    # Parse the XML
    # root = etree.fromstring(xml_string)
    # elem = root.find(".//commentReference")
    # if elem:
    #     log.debug(f"found comment! {elem}")
    #     return elem.get("id", None)  # Default to empty string
    # That's not working, try regex
    # <w:commentReference w:id="3">
    # HACK but it works
    regex = re.compile(r"w:commentReference w:id=\"(\d+)\"")
    m = regex.search(xml_string)
    if m:
        return int(m[1])
    else:
        return None



def extract_attribute_safely(tree, xpath, attr):
    """Extract attribute with proper None checking"""
    element = tree.find(xpath)
    if element is not None:
        return element.get(attr, "")  # Default to empty string
    return ""



def save_image(image_part, output_folder):
    """Save an image to the output folder and return the filename."""
    os.makedirs(output_folder, exist_ok=True)
    image_filename = os.path.join(output_folder, os.path.basename(image_part.partname))
    with open(image_filename, "wb") as img_file:
        img_file.write(image_part.blob)
    return str(image_filename).replace("\\", "/")


def get_list_level(paragraph):
    """Determine the level of a bullet point or numbered list item."""
    # Access the raw XML of the paragraph
    p = paragraph._element
    numPr = p.find(".//w:numPr", namespaces=p.nsmap)
    if numPr is not None:
        ilvl = numPr.find(".//w:ilvl", namespaces=p.nsmap)
        if ilvl is not None:
            return int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))
    return 0


def get_list_marker(paragraph):
    p = paragraph._element
    numPr = p.find(".//w:numPr", namespaces=p.nsmap)
    ilvl = numPr.find(".//w:numId", namespaces=p.nsmap)
    type_id = int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))

    match type_id:
        case 1:
            return '1. '
        case 2:
            return '- [ ] '
        case 3:
            return '* '
        case _:
            log.debug(f"Unknown list type id: {type_id}")
            # got for 9, 5, 8, 10.
            # need a way of looking up.

    # by default
    return "* "


def is_list(paragraph) -> bool:
    p = paragraph._element
    numPr = p.find(".//w:numPr", namespaces=p.nsmap)
    return numPr is not None


def get_bullet_point_prefix(paragraph):
    """
    Determine the Markdown prefix for a bullet point
    based on its indentation level.
    """
    level = get_list_level(paragraph)
    marker = get_list_marker(paragraph)
    return "  " * level + marker


def parse_run(run):
    """Go through document objects recursively and return markdown."""
    sub_parts = list(run.iter_inner_content())
    text = ""
    for s in sub_parts:
        if isinstance(s, str):
            text += s
        elif isinstance(s, docx.text.run.Run):
            text += parse_run(s)
        elif isinstance(s, docx.text.hyperlink.Hyperlink):
            text += f"[{s.text}]({s.address})"
        elif isinstance(s, docx.drawing.Drawing):
            rId = extract_r_embed(s._element.xml)
            text += f"![][image{rId[3:]}]"
        else:
            log.warning("unknown run type", s)

    if isinstance(run, docx.text.run.Run):
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        if run.underline:
            text = f"__{text}__"
        if run.font.strike:
            text = f"~~{text}~~"
        # check .font for monospacing
        # check style
        if run.font.name == "Courier New": # more fonts!
            text = f"`{text}`<br>"

        comment_id = extract_comment_id(run._element.xml)
        if comment_id:
            #log.debug(f"### found comment ID: {comment_id}")
            text += f"[^{comment_id}]"

# <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
# xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
# xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
# xmlns:w10="urn:schemas-microsoft-com:office:word"
# xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
# xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
# xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
# xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
# xmlns:v="urn:schemas-microsoft-com:vml"
# xmlns:o="urn:schemas-microsoft-com:office:office"
# xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
#  <w:commentReference w:id="3"/>
#</w:r>

    #else:
    #
    #     log.warning(f"type: {type(run)}, {run}")

    return text


class DocxitConverter(Converter):
    def convert(self, infile:Path, outroot:Path) -> Page:
        """
        Converts a docx file into markdown using docxit
        """
        return convert_out(infile, self.root, outroot)


    @staticmethod
    def load_file(full_path:Path) -> Page:
        """
        Given an DOCX file, load the content and convert to MD using
        the Docxit converter. This generates an in-memory Page object
        for the document with all attachments embedded.
        Given an MD file, just load a page with it.
        """
        match full_path.suffix.lower():
            case ".docx":
                return convert(full_path)
            case ".md":
                return Page.load_file(full_path)
            case _:
                log.info(f"Skipping, unknown file extension: {full_path}")
