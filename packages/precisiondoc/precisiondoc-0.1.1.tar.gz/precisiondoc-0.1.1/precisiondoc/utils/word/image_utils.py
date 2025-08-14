"""
Image utilities for Word documents.
"""
import os
from docx.shared import Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from precisiondoc.config.document_format_settings import IMAGE_SETTINGS

class ImageUtils:
    """Handles image-related operations"""
    
    @staticmethod
    def _add_image_to_cell(cell, image_path, image_col_width):
        """
        Add image to table cell with proper sizing
        
        Args:
            cell: Table cell to add image to
            image_path: Path to image file
            image_col_width: Width of the column in inches
            
        Returns:
            bool: True if image was added successfully, False otherwise
        """
        # Clear any existing content in the cell
        for paragraph in cell.paragraphs:
            p = paragraph._p
            p.getparent().remove(p)
        
        # Create a new paragraph for the image
        paragraph = cell.add_paragraph()
        paragraph.alignment = 1  # Center alignment
        
        # Set cell vertical alignment
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Check if image exists
        if not image_path or not os.path.exists(image_path):
            # Add placeholder text if image not found
            run = paragraph.add_run("Image not available")
            run.italic = True
            return False
        
        try:
            # Calculate image width (slightly smaller than column width)
            max_width = Inches(image_col_width.inches * IMAGE_SETTINGS['width_ratio'])
            
            # Add the image
            run = paragraph.add_run()
            run.add_picture(image_path, width=max_width)
            return True
        except Exception as e:
            # Add error text if image loading fails
            run = paragraph.add_run(f"Error loading image: {str(e)}")
            run.italic = True
            return False
