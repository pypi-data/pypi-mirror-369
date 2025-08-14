"""
Content formatting utilities for Word documents.
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from precisiondoc.config.document_format_settings import WORD_FONT_SETTINGS

class ContentFormatter:
    """Handles content-level formatting operations"""
    
    @staticmethod
    def apply_paragraph_format(paragraph):
        """
        Apply paragraph formatting
        
        Args:
            paragraph: Paragraph object
        """
        p_format = paragraph.paragraph_format
        # set line spacing to single
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Single line spacing
        # # set line spacing to 1.5
        # p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5 line spacing
        # # set line spacing to 1.3
        # p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        # p_format.line_spacing = 1.3  # 1.3 line spacing
        p_format.space_before = Pt(12)  # Space before paragraph
        p_format.space_after = Pt(12)   # Space after paragraph
    
    @staticmethod
    def apply_run_format(run):
        """
        Apply formatting to text run object
        
        Args:
            run: Text run object
        """
        run.font.name = WORD_FONT_SETTINGS['run_font_name']  # Set font to Microsoft YaHei
        run.font.size = Pt(WORD_FONT_SETTINGS['run_font_size'])     # Set font size
    
    @staticmethod
    def apply_separator_format(paragraph):
        """
        Apply formatting to separator paragraph
        
        Args:
            paragraph: Paragraph object
        """
        separator_format = paragraph.paragraph_format
        separator_format.space_before = Pt(12)
        separator_format.space_after = Pt(12)
        
        # Set separator line font
        for run in paragraph.runs:
            run.font.name = WORD_FONT_SETTINGS['separator_font_name']
    
    @staticmethod
    def add_text_to_cell(cell, text, multi_line=True):
        """
        Add text to a table cell, either as a single paragraph or multiple paragraphs
        
        Args:
            cell: Table cell object
            text: Text to add
            multi_line: If True, split text by newlines and create a paragraph for each line
        """
        # Clear default paragraph if it exists and is empty
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if not p.runs:  # Only clear if it's empty
                p._element.getparent().remove(p._element)
                p._p = None
                p._element = None
        
        if multi_line:
            # Split text by lines and create a paragraph for each line
            for line in text.strip().split('\n'):
                if line:
                    p = cell.add_paragraph()
                    p.alignment = 0  # Left alignment
                    run = p.add_run(line)
                    ContentFormatter.apply_run_format(run)
        else:
            # Add all text as a single paragraph
            p = cell.add_paragraph()
            p.alignment = 0  # Left alignment
            run = p.add_run(text)
            ContentFormatter.apply_run_format(run)
