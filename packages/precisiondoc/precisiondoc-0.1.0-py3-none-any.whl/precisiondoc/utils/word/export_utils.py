"""
Export utilities for Word documents.
"""
import os
import pandas as pd
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from precisiondoc.utils.word.document_formatting import DocumentFormatter
from precisiondoc.utils.word.evidence_processing import EvidenceProcessor
from precisiondoc.utils.word.content_formatting import ContentFormatter
from precisiondoc.utils.word.table_utils import TableUtils
from precisiondoc.utils.word.image_utils import ImageUtils

class ExportUtils:
    """Handles document export operations"""
    
    @staticmethod
    def _create_evidence_table(doc, evidence_dict, image_path, multi_line_text=True, show_borders=True, exclude_columns=None):
        """
        Create a table for evidence with text and image placeholders
        
        Args:
            doc: Word document
            evidence_dict: Dictionary containing evidence data
            image_path: Path to image file
            multi_line_text: Whether to use multi-line text format (one row per key-value pair)
            show_borders: Whether to show table borders
            exclude_columns: Columns to exclude from evidence text
            
        Returns:
            table: Created table
        """
        if exclude_columns is None:
            exclude_columns = ['is_precision_evidence', 'page_number', 'document_type']
        
        # Rename text to resource_sentence, image_path to resource_url (key name replacement)
        evidence_dict = EvidenceProcessor._rename_evidence_dict_keys(evidence_dict)
        
        # Sort the evidence dictionary based on the predefined key sequence
        evidence_dict = EvidenceProcessor._sort_evidence_dict(evidence_dict)
        
        # Create table based on format (multi-line or single-row JSON)
        if multi_line_text:
            # Count non-excluded items with non-empty values
            row_count = sum(1 for k, v in evidence_dict.items() 
                           if k not in exclude_columns and v and str(v).strip())
        else:
            row_count = 1
        
        # Create basic table
        table = TableUtils._create_basic_table(doc, row_count)
        
        # Get current section and check if it's landscape
        current_section = doc.sections[-1]
        is_landscape = current_section.orientation == 1  # 1 is landscape
        
        # Set table column widths
        table_width, text_col_width, image_col_width = TableUtils._set_table_column_widths(table, is_landscape)
        
        # Set table pagination properties
        TableUtils._set_table_pagination_properties(table)
        
        # Populate table cells
        if multi_line_text:
            # Populate with one key-value pair per row
            left_cells, right_cell = EvidenceProcessor._populate_multi_line_cells(table, evidence_dict, exclude_columns)
        else:
            # Populate with JSON-style dictionary in a single cell
            left_cell, right_cell = EvidenceProcessor._populate_json_format_cell(table, evidence_dict)
        
        # Add image to the right column
        ImageUtils._add_image_to_cell(right_cell, image_path, image_col_width)
        
        # Remove table borders if specified
        if not show_borders:
            TableUtils._remove_table_borders(table)
            
        return table
    
    @staticmethod
    def _add_separator_and_section(doc):
        """
        Add a separator line and section break
        
        Args:
            doc: Word document
        """
        # Add separator line
        separator = doc.add_paragraph('─' * 100)
        ContentFormatter.apply_separator_format(separator)
        
        # Add section break for next evidence
        doc.add_section()
    
    @staticmethod
    def export_evidence_to_word(excel_file, word_file=None, output_folder=None, multi_line_text=True, show_borders=True, exclude_columns=None):
        """
        Export precision evidence from Excel to Word document
        
        Args:
            excel_file: Path to Excel file or DataFrame
            word_file: Path to output Word file
            output_folder: Output folder path, used to find images
            multi_line_text: If True, split text by newlines in the left cell
            show_borders: If True, show table borders
            exclude_columns: Columns to exclude from evidence text
        """
        logger = logging.getLogger(__name__)
        
        try:
            # Check if excel_file is already a DataFrame
            if isinstance(excel_file, pd.DataFrame):
                df = excel_file
            else:
                # Read Excel file into DataFrame
                df = pd.read_excel(excel_file)
            
            # If word_file is None, generate a default output path
            if word_file is None:
                # Create output directory in current working directory if it doesn't exist
                output_dir = os.path.join(os.getcwd(), "output", "word")
                
                if isinstance(excel_file, str):
                    # Get the base name of the Excel file and change extension to .docx
                    base_name = os.path.basename(excel_file)
                    base_name_no_ext = os.path.splitext(base_name)[0]
                    word_file = os.path.join(output_dir, f"{base_name_no_ext}.docx")
                else:
                    # If excel_file is a DataFrame, use a timestamp-based name
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    word_file = os.path.join(output_dir, f"evidence_export_{timestamp}.docx")
                
                logger.info(f"No output file specified, using default: {word_file}")
            
            # Check if 'is_precision_evidence' column exists
            if 'is_precision_evidence' in df.columns:
                # Filter for precision evidence
                true_values = [True, 'True', 'true', 1, '1', 'yes', 'Y']
                evidence_df = df[df['is_precision_evidence'].isin(true_values)].copy()
                
                if evidence_df.empty:
                    logger.warning("No precision evidence found in Excel file")
                    return word_file
                
                # Create Word document
                doc = Document()
                
                # Apply document formatting
                DocumentFormatter.apply_word_document_format(doc)
                
                # Define default columns to exclude from evidence text if not provided
                if exclude_columns is None:
                    exclude_columns = (
                        'page_type', 'is_precision_evidence',
                        'page_number', 'success', 'document_type', '解析', '分析', '结论', '文字提取', 
                        'evidence_level', 'evidence_type', 'evidence_list'
                    )
                
                # Process evidence row by row
                for idx, row in evidence_df.iterrows():
                    # Create a new section for each evidence item to control page layout
                    if idx > 0:
                        new_section = doc.add_section()
                        # By default, use landscape orientation
                        DocumentFormatter.set_section_orientation(new_section, 'landscape')
                    
                    # Prepare evidence text
                    evidence_dict = EvidenceProcessor._prepare_evidence_text(row, exclude_columns)
                    
                    # Get image path
                    image_path = row.get('image_path', '')
                    
                    # Create table with text and image placeholders
                    ExportUtils._create_evidence_table(doc, evidence_dict, image_path, 
                                                      multi_line_text=multi_line_text,
                                                      show_borders=show_borders,
                                                      exclude_columns=exclude_columns)
                    
                    # Add separator line
                    separator = doc.add_paragraph('-' * 50)
                    ContentFormatter.apply_separator_format(separator)
                
                # Create directory if it doesn't exist
                if word_file is not None:
                    os.makedirs(os.path.dirname(word_file), exist_ok=True)
                    
                    # Save Word document
                    doc.save(word_file)
                    logger.info(f"Evidence exported to Word file: {word_file}")
                else:
                    logger.error("No output file path provided and could not generate a default path")
                return word_file
            else:
                logger.warning("No 'is_precision_evidence' column found in Excel file")
                return word_file
        except Exception as e:
            logger.error(f"Error exporting evidence to Word: {str(e)}")
            raise
