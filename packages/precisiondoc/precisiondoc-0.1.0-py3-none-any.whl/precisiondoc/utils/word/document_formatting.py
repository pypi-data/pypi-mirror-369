"""
Document formatting utilities for Word documents.
"""
import os
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from precisiondoc.config.document_format_settings import WORD_FONT_SETTINGS

class DocumentFormatter:
    """Handles document-level formatting operations"""
    
    @staticmethod
    def apply_word_document_format(doc):
        """
        Apply uniform formatting to Word document
        
        Args:
            doc: Word document object
        """
        # Set page margins and width
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(8.27)  # A4 width
            section.page_height = Inches(11.69)  # A4 height
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            # Default to portrait orientation
            # section.orientation = WD_ORIENT.PORTRAIT  # or WD_ORIENT.LANDSCAPE
            section.orientation = WD_ORIENT.LANDSCAPE

        # Add heading
        heading = doc.add_heading('Precision Evidence Report', 0)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(24)
        heading_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Set heading font
        for run in heading.runs:
            run.font.name = WORD_FONT_SETTINGS['heading_font_name']
            run.font.size = Pt(WORD_FONT_SETTINGS['heading_font_size'])
            run.font.bold = True
        
        # Add page numbers
        DocumentFormatter.add_page_numbers(doc)
    
    @staticmethod
    def add_page_numbers(doc):
        """
        Add page numbers in 'current/total' format to the document footer
        
        Args:
            doc: Word document object
        """
        # Add page numbers to each section
        for section in doc.sections:
            footer = section.footer
            
            # Clear any existing content in the footer
            for paragraph in footer.paragraphs:
                p = paragraph._p
                p.getparent().remove(p)
            
            # Add a paragraph to the footer
            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page number field codes
            run = paragraph.add_run()
            
            # Add the 'current page' field
            fld_char1 = OxmlElement('w:fldChar')
            fld_char1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_char1)
            
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = ' PAGE '
            run._r.append(instr_text)
            
            fld_char2 = OxmlElement('w:fldChar')
            fld_char2.set(qn('w:fldCharType'), 'end')
            run._r.append(fld_char2)
            
            # Add the separator
            run.add_text(' / ')
            
            # Add the 'total pages' field
            fld_char3 = OxmlElement('w:fldChar')
            fld_char3.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_char3)
            
            instr_text2 = OxmlElement('w:instrText')
            instr_text2.set(qn('xml:space'), 'preserve')
            instr_text2.text = ' NUMPAGES '
            run._r.append(instr_text2)
            
            fld_char4 = OxmlElement('w:fldChar')
            fld_char4.set(qn('w:fldCharType'), 'end')
            run._r.append(fld_char4)
            
            # Apply formatting to the page number text
            run.font.size = Pt(10)
            run.font.name = WORD_FONT_SETTINGS['page_number_font_name']
    
    @staticmethod
    def set_section_orientation(section, orientation='portrait'):
        """
        Set the orientation for a specific section
        
        Args:
            section: Section object
            orientation: 'portrait' or 'landscape'
        """
        if orientation.lower() == 'portrait':
            section.orientation = WD_ORIENT.PORTRAIT
            # When changing orientation, we need to swap width and height
            section.page_width = Inches(8.27)  # A4 width
            section.page_height = Inches(11.69)  # A4 height
        elif orientation.lower() == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            # When changing orientation, we need to swap width and height
            section.page_width = Inches(11.69)  # A4 height becomes width
            section.page_height = Inches(8.27)  # A4 width becomes height
