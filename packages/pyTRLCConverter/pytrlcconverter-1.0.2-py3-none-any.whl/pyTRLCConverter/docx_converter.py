"""Converter to Word docx format.

    Author: Norbert Schulz (norbert.schulz@newtec.de)
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2025 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************
import os
from typing import Optional
import docx
from pyTRLCConverter.base_converter import BaseConverter
from pyTRLCConverter.logger import log_verbose
from pyTRLCConverter.ret import Ret
from pyTRLCConverter.trlc_helper import Record_Object

# Variables ********************************************************************

# Classes **********************************************************************

class DocxConverter(BaseConverter):
    """
    Converter to docx format.
    """

    OUTPUT_FILE_NAME_DEFAULT = "output.docx"

    def __init__(self, args: any) -> None:
        # lobster-trace: SwRequirements.sw_req_no_prj_spec
        # lobster-trace: SwRequirements.sw_req_docx
        # lobster-trace: SwRequirements.sw_req_docx_template
        """
        Initialize the docx converter.

        Args:
            args (any): The parsed program arguments.
        """
        super().__init__(args)

        if args.template is not None:
            log_verbose(f"Loading template file {args.template}.")

        self._docx = docx.Document(docx=args.template)

        # Ensure default table style is present in the document.
        if not 'Table Grid' in self._docx.styles:
            self._docx.styles.add_style('Table Grid', docx.enum.style.WD_STYLE_TYPE.TABLE, builtin=True)

    @staticmethod
    def get_subcommand() -> str:
        # lobster-trace: SwRequirements.sw_req_docx
        """ Return subcommand token for this converter.

        Returns:
            Ret: Status
        """
        return "docx"

    @staticmethod
    def get_description() -> str:
        # lobster-trace: SwRequirements.sw_req_docx
        """ Return converter description.
 
        Returns:
            Ret: Status
        """
        return "Convert into docx format."

    @classmethod
    def register(cls, args_parser: any) -> None:
        # lobster-trace: SwRequirements.sw_req_docx
        """Register converter specific argument parser.

        Args:
            args_parser (any): Argument parser
        """
        super().register(args_parser)

        BaseConverter._parser.add_argument(
            "-t",
            "--template",
            type=str,
            default=None,
            required=False,
            help="Load the given docx file as a template to append to."
        )
        BaseConverter._parser.add_argument(
            "-n",
            "--name",
            type=str,
            default=DocxConverter.OUTPUT_FILE_NAME_DEFAULT,
            required=False,
            help="Name of the generated output file inside the output folder " \
                f"(default = {DocxConverter.OUTPUT_FILE_NAME_DEFAULT})."
        )

    def convert_section(self, section: str, level: int) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_section
        """Process the given section item.

        Args:
            section (str): The section name
            level (int): The section indentation level

        Returns:
            Ret: Status
        """
        self._docx.add_heading(section, level)

        return Ret.OK

    def convert_record_object_generic(self, record: Record_Object, level: int, translation: Optional[dict]) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_record
        """
        Process the given record object in a generic way.

        The handler is called by the base converter if no specific handler is
        defined for the record type.

        Args:
            record (Record_Object): The record object.
            level (int): The record level.
            translation (Optional[dict]): Translation dictionary for the record object.
                                            If None, no translation is applied.

        
        Returns:
            Ret: Status
        """
        return self._convert_record_object(record, level, translation)

    def finish(self) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_file
        """Finish the conversion.

        Returns:
            Ret: Status
        """
        result = Ret.ERROR

        if self._docx is not None:
            output_file_name = self._args.name
            if 0 < len(self._args.out):
                output_file_name = os.path.join(self._args.out, self._args.name)

            log_verbose(f"Writing docx {output_file_name}.")
            self._docx.save(output_file_name)
            self._docx = None
            result = Ret.OK

        return result

    def _convert_record_object(self, record: Record_Object, level: int, translation: Optional[dict]) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_record
        """
        Process the given record object.

        Args:
            record (Record_Object): The record object.
            level (int): The record level.
            translation (Optional[dict]): Translation dictionary for the record object.
                                            If None, no translation is applied.

        Returns:
            Ret: Status
        """
        self._docx.add_heading(f"{record.name} ({record.n_typ.name})", level + 1)
        attributes = record.to_python_dict()

        table = self._docx.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True

        # Set table headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Element"
        header_cells[1].text = "Value"

        # Populate table with attribute key-value pairs
        for key, value in attributes.items():
            if value is None:
                value = self._empty_attribute_value

            if translation is not None:
                if key in translation:
                    key = translation[key]

            cells = table.add_row().cells
            cells[0].text = key
            cells[1].text = str(value)

        # Add a paragraph with the record object location
        p = self._docx.add_paragraph()
        p.add_run(f"from {record.location.file_name}:{record.location.line_no}").italic = True

        return Ret.OK


# Functions ********************************************************************

# Main *************************************************************************
