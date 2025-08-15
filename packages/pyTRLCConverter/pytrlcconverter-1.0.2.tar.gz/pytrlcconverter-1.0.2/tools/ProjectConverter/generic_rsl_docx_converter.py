"""Project specific docx converter for Generic.rsl types.

This module provides a project specific docx converter subclass with
support for the TRLC record types defined in Generic.rsl.

    Author: Norbert Schulz (norbert.schulz@newtec.de)
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2025 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************
from typing import Optional
import docx
# pylint: disable=import-error
from pyTRLCConverter.docx_converter import DocxConverter
from pyTRLCConverter.ret import Ret
from pyTRLCConverter.trlc_helper import Record_Object

# pylint: disable=wrong-import-order
from image_processing import convert_plantuml_to_image, locate_file

# Variables ********************************************************************

# Classes **********************************************************************

class GenericRslDocxConverter(DocxConverter):
    """
    Project specific docx converter subclass for generic.rsl types.
    """
    def __init__(self, args: any) -> None:
        """
        Initialize the custom docx converter.

        Args:
            args (any): The parsed program arguments.
        """
        super().__init__(args)

        self._img_counter = 1

    # pylint: disable=unused-argument
    def _convert_record_object_info(self, record: Record_Object, level: int, translation: Optional[dict]) -> Ret:
        """Convert an information record object to the destination format.

        Args:
            record (Record_Object): The record object to convert.
            level (int): Current level of the record object.
            translation (Optional[dict]): Translation dictionary for the record object.
                                            If None, no translation is applied.

        Returns:
            Ret: Status
        """
        self._docx.add_paragraph(self._get_attribute(record, "description"))
        return Ret.OK

    def _convert_record_object_plantuml(self, record: Record_Object, level: int, translation: Optional[dict]) -> Ret:
        """Convert a Plantuml diagram record object to the destination format.

        Args:
            record (Record_Object): The record object to convert.
            level (int): Current level of the record object.
            translation (Optional[dict]): Translation dictionary for the record object.
                                            If None, no translation is applied.

        Returns:
            Ret: Status
        """
        result = Ret.ERROR

        image_file = convert_plantuml_to_image(
            self._get_attribute(record, "file_path"),
            self._args.out,
            self._args.source
        )

        if image_file is not None:
            self._add_image(image_file, self._get_attribute(record, "caption"), level)
            result = Ret.OK

        return result

    def _convert_record_object_image(self, record: Record_Object, level: int, translation: Optional[dict]) -> Ret:
        """Convert a software diagram record object to the destination format.

        Args:
            record (Record_Object): The record object to convert.
            level (int): Current level of the record object.
            translation (Optional[dict]): Translation dictionary for the record object.
                                            If None, no translation is applied.

        Returns:
            Ret: Status
        """
        result = Ret.ERROR

        image_file = locate_file(self._get_attribute(record, "file_path"), self._args.source)
        if image_file is not None:
            self._add_image(image_file, self._get_attribute(record, "caption"), level)
            result = Ret.OK

        return result

    def _add_image(self, image_file: str, caption: str, level: int) -> None:
        """Add an image to the docx file.

        Args:
            image_file (str): The image file to add.
            caption (str): The caption of the image.
            level (int): Current level of the record object
        """
        p = self._docx.add_paragraph()
        p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_file, width=docx.shared.Inches(6))
        run.add_text(f"Figure {self._img_counter} {caption}")

        self._img_counter += 1

# Functions ********************************************************************

# Main *************************************************************************
