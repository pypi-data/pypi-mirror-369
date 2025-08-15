"""Test the docx requirements.
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2025 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************
import os
from argparse import Namespace
from collections import namedtuple
import docx

from pyTRLCConverter.__main__ import main
from pyTRLCConverter.docx_converter import DocxConverter

# Variables ********************************************************************

# Classes **********************************************************************

# Functions ********************************************************************

def test_tc_docx(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx
    """
    The software shall support conversion of TRLC source files into docx.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx")

    # Mock program arguments to simulate running the script with inbuild Markdown converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_no_section.trlc",
        "--out", str(tmp_path),
        "docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the requirement is present.
    assert created_docx.paragraphs[0].text == "req_id_1 (Requirement)"

    # Check the created table for the requirement.
    assert len(created_docx.tables) == 1
    assert created_docx.tables[0].cell(0, 0).text == "Element"
    assert created_docx.tables[0].cell(0, 1).text == "Value"
    assert created_docx.tables[0].cell(1, 0).text == "description"
    assert created_docx.tables[0].cell(1, 1).text == "Test description"

def test_tc_docx_section(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_section
    """
    The test case checks if a TRLC section is correctly converted into docx.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_section")

    # Mock program arguments to simulate running the script with inbuild Markdown converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_with_section.trlc",
        "--out", str(tmp_path),
        "docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the section is present
    assert created_docx.paragraphs[0].text == "Test section"

    # Check that the requirement is present.
    assert created_docx.paragraphs[1].text == "req_id_2 (Requirement)"

def test_tc_docx_file(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_file
    """
    The test case checks whether multiple TRLC files are combined into one docx output document.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_file")

    # Mock program arguments to simulate running the script with inbuild Markdown converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_no_section.trlc",
        "--source", "./tests/utils/single_req_with_section.trlc",
        "--out", str(tmp_path),
        "docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the first requirement is present.
    assert created_docx.paragraphs[0].text == "req_id_1 (Requirement)"
    # Paragraph 1 is the "from file" paragraph.
    # Check the secont requirement is present.
    assert created_docx.paragraphs[2].text == "Test section"
    assert created_docx.paragraphs[3].text == "req_id_2 (Requirement)"

def test_tc_docx_template(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_template
    """
    The test case checks whether docx templated are correctly applied.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_template")

    # Mock program arguments to simulate running the script with inbuild Markdown converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_no_section.trlc",
        "--out", str(tmp_path),
        "docx",
        "--template", "./tests/utils/template.docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the template was used.
    assert created_docx.paragraphs[0].text == "Template text."

    # Check that the requirement is present.
    assert created_docx.paragraphs[1].text == "req_id_1 (Requirement)"
